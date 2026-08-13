#!/usr/bin/env python3
"""生成《拟公示内容》《算法安全自评估报告》填写稿（docx）——按 2026-08 驳回意见整改版。

统一口径（全材料必须一致）：
- 算法名称：灵祺智能生成合成算法
- 备案主体：宁波墨典网络科技有限公司（与域名 ICP 主体一致）
- ICP：浙ICP备2026044830号-1（mofangdianai.com）
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT_DIR = Path(__file__).resolve().parent
DOWNLOADS = Path.home() / "Downloads"

# 与 ICP / 用户协议主体对齐；信用代码请与营业执照核对后替换
COMPANY = {
    "name": "宁波墨典网络科技有限公司",
    "credit_code": "【请填写统一社会信用代码，须与营业执照一致】",
    "legal_rep": "曹鑫淼",
    "algo_officer": "曹鑫淼",
    "phone": "15757468650",
    "address": "浙江省宁波市",
    "domain": "mofangdianai.com",
    "icp": "浙ICP备2026044830号-1",
    "products": "灵祺AI智能ERP（网站）、灵祺星选（微信小程序）及相关配套 Web",
    "product_urls": (
        "商家ERP网站：https://cs.mofangdianai.com；"
        "服务商网站：https://fws.mofangdianai.com；"
        "履约网站：https://dr.mofangdianai.com；"
        "微信小程序：灵祺星选；"
        "API：https://mofangdianai.com/erp-api"
    ),
}

ALGO = {
    "name": "灵祺智能生成合成算法",
    "type": "生成合成类",
    "field": "用户内容生成(UGC)、短视频、商超团购",
    "scene": (
        "本地生活商家营销内容辅助生成（经营文案、组品方案、活动话术等）；"
        "达人/商单相关说明与脚本辅助生成；"
        "在合规前提下提供图像、语音、短视频等生成合成能力调用与编排"
        "（含云剪成片、数字人口播等）。"
    ),
    "online_date": "2025年6月1日",
    "version": "V1.0",
    # 与填报项「输出模态 / 文件格式」严格对应
    "output_modality": "文本、音频、视频",
    "output_formats": "txt、json、mp3、mp4",
    "input_modality": "文本、图像、视频",
    "display_form": "网站、微信小程序",
}

# 第三方模型（提交前请用厂商公示页截图核对最新编号）
MODELS = [
    {
        "name": "豆包大模型（火山方舟 API）",
        "provider": "北京火山引擎科技有限公司等相关主体",
        "algo_beian": "网信算备110108823483901230031号（豆包大模型算法）",
        "model_beian": "Beijing-YunQue-20230821（以厂商公示为准）",
        "use": "文本对话、经营建议、组品方案、招募 Brief、文案生成等",
    },
    {
        "name": "通义千问大模型（阿里云百炼/DashScope API）",
        "provider": "阿里巴巴达摩院（杭州）科技有限公司等相关主体",
        "algo_beian": "网信算备330110507206401230035号（达摩院交互式多能型合成算法，以厂商公示为准）",
        "model_beian": "ZheJiang-TongYiQianWen-20230901（以厂商公示为准）",
        "use": "文本对话、内容生成、多模态辅助等",
    },
    {
        "name": "阿里云智能媒体服务 ICE / IMS（云剪辑、数字人、TTS 等）",
        "provider": "阿里云计算有限公司",
        "algo_beian": "按所用具体能力对应厂商已公示算法/模型备案号填写（提交前截图核对）",
        "model_beian": "按合同与控制台开通能力对应填写",
        "use": "短视频云剪成片、数字人口播、语音合成（mp3/mp4）",
    },
]

DEPLOY = (
    "软硬件均部署于中华人民共和国境内，不跨境存储与推理。"
    "（1）业务 API 与数据库：阿里云轻量应用服务器，公网 IP 139.196.42.5，"
    "机房位于阿里云中国内地华东地域（杭州节点所属区域），运行 Node.js 业务服务与 PostgreSQL 等；"
    "（2）静态 Web 前端：阿里云 ECS，公网 IP 8.160.173.236，位于阿里云中国内地华东地域，"
    "托管商家/服务商/履约网站静态资源与 Nginx；"
    "（3）对象存储与媒体合成：调用阿里云境内对象存储及智能媒体服务（ICE/IMS），数据驻留中国内地；"
    "（4）大模型推理：经 HTTPS 调用火山引擎、阿里云等境内已备案模型 API，不向境外传输训练用途数据。"
    f"对外服务域名：{COMPANY['domain']}（ICP：{COMPANY['icp']}）。"
)

TODAY = date.today().strftime("%Y年%m月%d日")


def set_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def body(doc: Document, text: str, indent: bool = True) -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "宋体"
                    r.font.size = Pt(11)
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def build_publicity() -> Document:
    doc = Document()
    set_font(doc)
    h1(doc, "拟公示内容")
    doc.add_paragraph()
    kv_table(
        doc,
        [
            ("算法名称", ALGO["name"]),
            (
                "算法基本原理",
                "本算法基于深度学习与大语言模型等生成合成技术，结合规则编排与业务模板，"
                "根据用户输入的业务指令、门店/商品信息、营销目标及音视频素材等，"
                f"自动或辅助生成{ALGO['output_modality']}内容"
                f"（文件格式主要为{ALGO['output_formats']}）。"
                "底层调用经合法授权且已备案/登记的第三方通用大模型与云媒体合成服务，"
                "由本平台完成提示词组织、安全过滤、结果编排与业务落地；不自研训练底层基础大模型。",
            ),
            (
                "算法运行机制",
                f"用户在{COMPANY['products']}中发起生成请求→平台校验登录与权限→"
                "组装业务上下文与提示词→调用生成合成模型/云媒体服务→"
                "对输出进行敏感内容过滤与合规校验→按产品界面返回结果供用户预览、编辑后使用。"
                "关键链路记录日志（账号、时间、任务标识、模型通道），支持人工客服复核与违规处置。"
                f"服务展现形态：{ALGO['display_form']}；访问路径见产品说明。",
            ),
            ("算法应用场景", ALGO["scene"]),
            (
                "算法目的意图",
                "提升本地生活商家与达人侧内容生产效率，辅助完成合规范围内的营销文案与素材生成，"
                "降低门店数字化运营成本；不用于编造新闻、不用于违法违规内容生产，不以误导公众为目的。",
            ),
            (
                "算法公示情况（选填）",
                f"本公示内容由{COMPANY['name']}依据《互联网信息服务算法推荐管理规定》"
                f"《互联网信息服务深度合成管理规定》编制，随算法备案材料一并提交。"
                f"产品与服务入口：{COMPANY['product_urls']}；"
                f"ICP备案号：{COMPANY['icp']}（主体与备案主体一致：{COMPANY['name']}）。",
            ),
        ],
    )
    doc.add_paragraph()
    body(doc, f"主体名称：{COMPANY['name']}", indent=False)
    body(doc, f"统一社会信用代码：{COMPANY['credit_code']}", indent=False)
    body(doc, f"算法安全负责人：{COMPANY['algo_officer']}", indent=False)
    body(doc, f"联系电话：{COMPANY['phone']}", indent=False)
    body(doc, f"公示日期：{TODAY}", indent=False)
    return doc


def build_self_assessment() -> Document:
    doc = Document()
    set_font(doc)
    h1(doc, "互联网信息服务算法安全自评估报告")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("（生成合成类-服务提供者）")
    r.font.size = Pt(12)

    doc.add_paragraph()
    kv_table(
        doc,
        [
            ("主体名称", COMPANY["name"]),
            ("统一社会信用代码", COMPANY["credit_code"]),
            ("算法名称", ALGO["name"]),
            ("算法类型", ALGO["type"]),
            ("算法应用领域", ALGO["field"]),
            ("算法使用场景", ALGO["scene"]),
            ("算法上线情况", f"☑已上线，时间：{ALGO['online_date']}"),
            ("自评估时间", TODAY),
            ("报告撰写时间", TODAY),
            (
                "算法基本情况",
                f"{ALGO['name']}面向本地生活商家与达人生态，提供{ALGO['output_modality']}等"
                f"生成合成能力编排服务（输出格式{ALGO['output_formats']}）。"
                f"算法作为{COMPANY['products']}的组成部分上线运行，展现形态为{ALGO['display_form']}，"
                f"由服务提供者直接向终端注册用户提供服务。版本{ALGO['version']}。"
                f"域名{COMPANY['domain']}，ICP：{COMPANY['icp']}。",
            ),
            ("算法备案类型", "☑算法未备案（本次首次申请备案）"),
            ("拟公示内容", f"“{ALGO['name']}”拟公示内容（附件）"),
            (
                "落实主体责任基本情况",
                f"“{COMPANY['name']}”落实算法安全主体责任基本情况（附件）",
            ),
            (
                "评估算法描述",
                "1.算法简介：基于大模型与生成合成技术，结合业务模板与安全策略，辅助生成营销与履约相关内容。\n"
                f"2.应用范围：{COMPANY['products']}中的 AI 生成功能（对话、组品、Brief、云剪、数字人等）。\n"
                "3.服务群体：本地生活商家、服务商、达人/PR 等注册用户（需登录）。\n"
                "4.用户数量：以实际运营统计为准（B端订阅与小程序注册用户）。\n"
                "5.社会影响：提高中小商家内容生产效率，影响范围限于平台注册用户业务场景，"
                "不面向社会大众信息流推荐，不提供新闻采编服务。\n"
                f"6.软硬件设施及部署位置：{DEPLOY}\n"
                "7.其他：底层模型通过合法授权的第三方 API 调用，平台负责业务编排与安全管控；"
                "详见「算法模型」及附件《第三方模型调用说明与证明材料》。",
            ),
            (
                "评估算法风险描述",
                "1.算法滥用风险：用户可能试图生成虚假宣传、侵权或违规内容。\n"
                "2.算法被恶意利用风险：批量调用、绕过审核、伪造身份生成不良内容。\n"
                "3.算法漏洞风险：提示注入、异常输入导致越权或错误输出。\n"
                "4.违法和不良信息生成风险：输出涉政、色情、暴力、歧视、谣言等。\n"
                "5.违法和不良信息存储风险：生成结果或日志被不当留存。\n"
                "6.违法和不良信息传播扩散风险：用户将违规内容外发至公域平台。\n"
                "7.数据和用户信息泄露风险：门店、用户个人信息被未授权访问或泄露。\n"
                "8.其他违法违规风险：未明示深度合成标识、未成年人保护不足等。",
            ),
            (
                "真实性声明",
                "我方承诺：提供的所有材料准确、真实、合法、有效，并愿为此承担有关法律责任。",
            ),
            ("算法安全负责人", f"{COMPANY['algo_officer']}（签名）"),
            ("联系电话", COMPANY["phone"]),
        ],
    )

    h2(doc, "一、算法情况")
    h3(doc, "（一）算法流程")
    body(
        doc,
        "流程节点：用户请求接入→身份与权限校验→业务上下文组装（门店/商品/任务描述）→"
        "提示词与参数生成→生成合成模型/云媒体服务调用（文本/图像/语音/视频等）→"
        "安全过滤与策略干预→结果封装返回并添加显式标识提示→（可选）人工客服或运营复核→"
        "用户确认后使用→日志留痕与溯源。",
    )
    body(
        doc,
        "节点粒度说明：身份校验与上下文组装为平台规则节点；模型调用为机器学习节点；"
        "安全过滤与内容审核为干预策略节点；人工复核为人工干预节点；标识与日志为合规节点。",
    )

    h3(doc, "（二）算法数据")
    body(
        doc,
        f"输入数据模态：{ALGO['input_modality']}。"
        "文本类输入：用户中文自然语言指令、门店名称、品类、活动卖点、目标平台等业务文本；"
        "一般不强制采集人脸等生物特征用于身份识别。涉及账号信息时仅用于鉴权与会话关联。",
    )
    body(
        doc,
        f"输出数据模态：{ALGO['output_modality']}；输出文件格式：{ALGO['output_formats']}。"
        "其中：文本结果以 txt/json（UTF-8）返回，体积通常为 KB 级；"
        "语音为 mp3；短视频/数字人成片为 mp4，单文件通常不超过 200MB。"
        "跨模态场景：文本→图像/语音/视频。",
    )
    body(
        doc,
        "图像/视频类输入（如启用）：用户上传的门店图片、参考图、菜单图或探店素材等；"
        "可能包含店铺环境画面；原则上不采集用于身份识别的生物特征。"
        "上传格式常见为 jpg/png/mp4，大小受产品限制。",
    )
    body(
        doc,
        "训练数据：本平台以调用经授权的通用大模型/生成服务为主，不在本地保存完整第三方模型训练语料，"
        "不自研训练底层基础大模型。如存在业务微调或提示词样本库，样本来源于公开语料、合成样例及"
        "经授权的业务脱敏数据，规模按内部版本管理，不包含违法违规内容。",
    )

    h3(doc, "（三）算法模型")
    body(
        doc,
        "1. 灵祺业务编排模型（平台侧自研）：负责将 ERP/小程序业务字段转换为生成请求，"
        "进行参数约束与结果结构化；类型为规则+模板+接口编排，随产品迭代更新；不构成底层大模型训练。",
    )
    body(doc, "2. 所调用第三方技术模型（能力层，须附合同/开通截图及厂商公示截图）：")
    for i, m in enumerate(MODELS, 1):
        body(
            doc,
            f"（{i}）模型/服务名称：{m['name']}；提供方：{m['provider']}；"
            f"相关算法备案号：{m['algo_beian']}；模型备案/上线信息：{m['model_beian']}；"
            f"在本产品中的用途：{m['use']}。",
        )
    body(
        doc,
        "模型类型主要为 Transformer 等深度学习结构及云媒体合成管线；"
        "优化目标为在给定提示下生成可用的业务文案或素材；"
        "评价指标包括可用性、违规命中率、延迟等；更新迭代跟随供应商版本与本平台接入策略。",
    )
    body(
        doc,
        "3. 内容安全检测策略：对输入输出进行敏感词与违规类别检测（规则库+云安全能力），"
        "命中则拦截或提示用户修改。",
    )
    body(
        doc,
        "说明：本备案算法作为服务提供者对外提供生成合成服务；"
        "不单独提供人脸替换类深度伪造娱乐产品；数字人口播使用商用预设形象/音色，"
        "不采集用户声纹用于克隆。如后续能力重大变更，将按规定办理变更备案。"
        "【附件】第三方模型合同/API 开通截图、厂商备案公示截图。",
    )

    h3(doc, "（四）干预策略（含输入/结果审核方式、审核范围）")
    body(
        doc,
        "【输入数据的审核方式】机器自动审核为主，覆盖全部在线生成接口："
        "（1）身份与权限校验（登录态、租户隔离、套餐配额）；"
        "（2）输入长度/格式校验与必要脱敏；"
        "（3）敏感词与违规意图预检（涉政、色情、暴恐、歧视、违禁品、虚假宣传等类别）；"
        "（4）禁止将合同底价等敏感字段上传至未评估的公共免费 AI。"
        "命中高风险输入时直接拒答或阻断请求。审核范围：全部注册用户发起的生成请求。",
    )
    body(
        doc,
        "【结果的审核方式】机器审核 + 人工抽检/客服复核："
        "（1）生成后对输出再次敏感词/违规类别过滤；"
        "（2）产品侧强制「AI 起草→预览→用户人工确认」后，方可写入业务库或用于对外发布链路；"
        "（3）用户举报、高风险账号及抽检样本由客服/运营人工复核，必要时暂停账号生成权限；"
        "（4）不提供新闻合成发布能力。"
        "审核范围：全部生成输出；写操作与对外发布相关结果实行人工确认门禁。"
        "上述机制用于证明安全风险防范的有效性：自动拦截降低违规产出，人工确认阻断自动外发。",
    )
    body(
        doc,
        "权限与配额干预：按租户套餐与账号权限控制调用；异常高频请求可限流或封禁。",
    )

    h3(doc, "（五）结果标识（须附截图证明）")
    body(
        doc,
        "1. 溯源标识（隐式）：平台侧对生成请求保留日志，至少包括账号标识、时间、会话/任务标识、"
        "功能模块、模型通道等，具备内部追踪溯源能力；音视频成片对接云厂商 AIGC 隐式标识/文件元数据能力"
        "（随 ICE/数字人服务侧能力启用）。【附件】后台日志字段截图或审计说明页截图。",
    )
    body(
        doc,
        "2. 显性标识：在产品界面以显著文字提示「AI生成 / 人工智能生成 / AI辅助生成，请人工确认」；"
        "文本结果展示区顶部或文首标注；视频/数字人成片在片头或画面角落标注；"
        "用户协议与隐私政策要求用户对外发布时依法保留显著标识。"
        "【附件】生成结果页显性标识截图、成片标注截图、用户协议相关条款截图。",
    )

    h2(doc, "二、服务情况（展示形态须与填报项产品及功能信息一致）")
    h3(doc, "（一）灵祺AI智能ERP（网站）")
    body(
        doc,
        f"服务简介：面向本地生活商家的 SaaS 工作台，提供门店经营、达人招募、"
        f"AI 营销文案/素材辅助生成、云剪/数字人等能力。"
        f"展现形态：网站（非 App）。访问地址：https://cs.mofangdianai.com。"
        f"上线时间：{ALGO['online_date']}起分阶段上线。"
        f"入口位于商家后台 AI 相关菜单与创作工具。用户为付费/注册商家，使用需登录（实名/手机号或微信授权）。"
        f"前置许可：ICP 备案号 {COMPANY['icp']}，备案主体为{COMPANY['name']}，与本算法备案主体一致。",
    )
    body(
        doc,
        "算法在服务中应用情况：数据来源于商家填写的门店与活动信息及用户指令；"
        "训练语料不落本地全量复制；模型能力按需调用境内云端接口；"
        "更新频率随产品发版与模型供应商升级；中间结果原则上不向无关第三方共享，"
        "仅用于完成本次生成与安全审计。",
    )
    h3(doc, "（二）灵祺星选（微信小程序）及配套履约网站")
    body(
        doc,
        "服务简介：面向达人、PR 与合作方的撮合与履约协同。"
        "展现形态：微信小程序（灵祺星选）+ 配套网站 https://dr.mofangdianai.com。"
        "在合规范围内提供话术、说明、脚本及云剪相关辅助生成。"
        "入口位于小程序及运营相关页面；使用需登录。"
        f"域名侧 ICP 同为 {COMPANY['icp']}（主体 {COMPANY['name']}）。",
    )
    body(
        doc,
        "算法在服务中应用情况：与 ERP 侧共用算法安全策略与日志机制；按账号权限开放生成能力。",
    )

    h2(doc, "三、风险研判")
    h3(doc, "（一）算法滥用")
    body(
        doc,
        "存在用户利用生成能力制作虚假促销、侵权或不良内容的风险。影响包括误导消费者、损害平台与商家信誉、"
        "引发投诉与监管风险。本平台通过协议禁止、机器审核、人工确认门禁与处置降低该风险。",
    )
    h3(doc, "（二）算法漏洞")
    body(
        doc,
        "存在提示注入、异常参数导致越权输出或不稳定结果的风险，可能影响业务连续性或产生不当内容。"
        "通过输入校验、权限隔离、超时与重试控制、模型输出复核降低风险。",
    )
    h3(doc, "（三）算法恶意利用")
    body(
        doc,
        "存在被第三方批量爬取接口、盗用密钥或诱导生成违禁内容的风险。通过鉴权、密钥管理、限流、"
        "异常监测与封禁降低风险。",
    )
    h3(doc, "其他风险")
    body(
        doc,
        "还包括个人信息泄露、生成内容外发传播、未成年人不当使用等风险，纳入内容治理与隐私合规统筹防控。",
    )

    h2(doc, "四、风险防控情况")
    h3(doc, "（一）风险防范机制建设")
    body(
        doc,
        "算法机制机理审核、算法安全评估监测、对生成合成虚假信息的辟谣机制、算法安全事件应急处置等，"
        f"详见附件《{COMPANY['name']}落实算法安全主体责任基本情况》。对第三章所列滥用、漏洞、恶意利用及"
        "违法不良信息风险均有效。",
    )

    h3(doc, "用户权益保护（须附截图证明）")
    body(
        doc,
        "用户知情权：通过用户协议、隐私政策及生成界面提示告知 AI 辅助生成情况；"
        "处理个人信息依法告知并取得同意（注册/登录与功能授权流程）。"
        "【附件】用户协议/隐私政策含生成式 AI 条款截图；注册登录同意页截图；生成页知情提示截图。",
    )
    body(
        doc,
        "用户个人信息保护：遵循《个人信息保护法》等要求，最小必要采集；生成业务数据与第三方模型服务交互时，"
        "仅传输完成生成所必需字段；不向无关第三方出售个人信息；共享遵循合同与审批。"
        "【附件】隐私政策个人信息处理条款截图；权限/租户隔离说明或后台权限截图。",
    )
    body(
        doc,
        "其他权益保护：同时遵循网络安全法、电子商务法、广告法、消费者权益保护法、未成年人保护法等；"
        "禁止生成违法广告与侵害他人权益内容；提供投诉举报与客服渠道（电话 "
        f"{COMPANY['phone']}）。【附件】投诉入口或客服联系方式截图。",
    )

    h3(doc, "内容生态治理（须附截图证明）")
    body(
        doc,
        "防范和抵制违法违规不良信息：建立敏感词与违规类别拦截、生成失败提示、用户举报与客服处置闭环；"
        "对违法不良信息生成、存储、传播风险有效。"
        "【附件】拦截/拒答效果截图；举报或处置流程截图（如有）。",
    )
    body(
        doc,
        "人工审核：对投诉举报、高风险账号及抽检样本开展人工复核；机器审核与人工审核相结合，"
        "写操作强制用户确认；必要时暂停相关账号生成权限。"
        "【附件】「预览确认后写入」交互截图。",
    )

    h3(doc, "模型安全保障")
    body(
        doc,
        "接口鉴权与密钥轮换保障机制：限制未授权调用；对异常流量告警。对恶意利用、滥用风险有效。"
        "提示与输出安全策略：拒绝明显违规请求；对输出二次过滤。对违法不良信息生成风险有效。",
    )

    h3(doc, "数据安全防护")
    body(
        doc,
        "训练/样本与业务数据采集遵循合法、正当、必要原则；存储于境内云资源，按权限访问；"
        "与第三方模型服务交互通过加密信道（HTTPS）；日志分级保存与定期审计，降低泄露风险。"
        f"部署位置详见评估算法描述第 6 点。",
    )

    h2(doc, "五、安全评估结论")
    body(
        doc,
        f"经自评估，{ALGO['name']}在现有业务场景下，已建立与主要风险相匹配的安全策略与处置机制"
        "（输入机器审核、输出机器审核+人工确认门禁、显式/溯源标识、用户权益与内容治理措施），"
        "具备上线运行条件。本报告结论为：风险总体可控，同意按服务提供者角色申请算法备案并持续改进安全措施。"
        "后续若算法能力、场景或主体信息发生重大变更，将按规定办理变更备案。",
    )

    h2(doc, "六、其他应当说明的相关情况")
    body(
        doc,
        "1. 本算法底层使用第三方云厂商大模型/生成 API 与云媒体服务，本单位作为服务提供者对面向用户的服务与安全负责；"
        "第三方模型名称、备案号及合同/开通证明见第三章及附件。\n"
        "2. 算法名称在填报项、《拟公示内容》与本报告中统一为「灵祺智能生成合成算法」，不得使用其他别名。\n"
        f"3. 输出模态（{ALGO['output_modality']}）与文件格式（{ALGO['output_formats']}）一一对应；"
        f"展示形态（{ALGO['display_form']}）与产品及功能信息一致。\n"
        f"4. ICP 前置许可填写 {COMPANY['icp']}，备案主体为{COMPANY['name']}，与本算法备案主体一致。\n"
        "5. 附件包括：拟公示内容；落实算法安全主体责任基本情况；算法备案承诺书；"
        "第三方模型合同/开通与公示截图；结果标识与用户权益/内容治理截图等。",
    )

    doc.add_paragraph()
    body(doc, f"主体名称（盖章）：{COMPANY['name']}", indent=False)
    body(doc, f"算法安全负责人（签名）：{COMPANY['algo_officer']}", indent=False)
    body(doc, f"日期：{TODAY}", indent=False)
    return doc


def main() -> None:
    publicity = build_publicity()
    assessment = build_self_assessment()

    names = [
        (publicity, f"拟公示内容-{ALGO['name']}-驳回整改版.docx"),
        (assessment, f"算法安全自评估报告-{ALGO['name']}-驳回整改版.docx"),
    ]
    for doc, name in names:
        for folder in (OUT_DIR, DOWNLOADS):
            folder.mkdir(parents=True, exist_ok=True)
            path = folder / name
            doc.save(path)
            print("wrote", path)


if __name__ == "__main__":
    main()
